from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


OUTPUT_DOCX = "Тренажер по номенклатуре - ответы.docx"

# num, formula, systematic, trivial (None = нет тривиального)
ANSWERS = [
    # Простейшие углеводороды
    (1,  "H3C–CH2–CH3",               "пропан",                   None),
    (2,  "H3C–CH2–CH2–CH3",           "бутан",                    "н-бутан"),
    (3,  "H2C=CH2",                   "этен",                     "этилен"),
    (4,  "HC≡CH",                     "этин",                     "ацетилен"),
    (5,  "C₆H₆",                      "бензол",                   None),
    # Спирты и фенолы
    (6,  "H3C–OH",                    "метанол",                  "метиловый спирт"),
    (7,  "H3C–CH2–OH",                "этанол",                   "этиловый спирт"),
    (8,  "H3C–CH(OH)–CH3",            "пропан-2-ол",              "изопропанол (изопропиловый спирт)"),
    (9,  "HO–CH2–CH2–OH",             "этан-1,2-диол",            "этиленгликоль"),
    (10, "C₆H₅–OH",                   "фенол (гидроксибензол)",   "карболовая кислота"),
    # Альдегиды и кетоны
    (11, "HC(=O)H",                   "метаналь",                 "формальдегид"),
    (12, "H3C–CH=O",                  "этаналь",                  "уксусный альдегид (ацетальдегид)"),
    (13, "H3C–C(=O)–CH3",             "пропан-2-он",              "ацетон"),
    (14, "H3C–CH2–CH=O",              "пропаналь",                "пропионовый альдегид"),
    # Карбоновые кислоты
    (15, "HCOOH",                     "метановая кислота",        "муравьиная кислота"),
    (16, "CH3COOH",                   "этановая кислота",         "уксусная кислота"),
    (17, "H3C–CH2–COOH",              "пропановая кислота",       "пропионовая кислота"),
    (18, "HOOC–COOH",                 "этандиовая кислота",       "щавелевая кислота"),
    # Амины и нитросоединения
    (19, "CH3–NH2",                   "метанамин",                "метиламин"),
    (20, "CH3–NO2",                   "нитрометан",               None),
    # Эфиры и галогенпроизводные
    (21, "CH3–O–CH3",                 "метоксиметан",             "диметиловый эфир"),
    (22, "H3C–C(=O)–O–CH3",           "метилэтаноат",             "метилацетат"),
    (23, "H3C–CH2–Cl",                "хлорэтан",                 "этилхлорид"),
    (24, "Cl–CH2–CH2–Cl",             "1,2-дихлорэтан",           "этилендихлорид"),
    # Разветвлённые цепи
    (25, "H3C–C(CH3)2–CH3",           "2,2-диметилпропан",        "неопентан"),
    (26, "H3C–CH(CH3)–CH2–CH3",       "2-метилбутан",             "изопентан"),
    (27, "H3C–C(CH3)=CH–CH3",         "2-метилбут-2-ен",          None),
    # Циклические соединения
    (28, "C₅H₁₀ (цикл)",              "циклопентан",              None),
    (29, "C₆H₅–CH3",                  "метилбензол",              "толуол"),
    (30, "H3C–CH(NH2)–COOH",          "2-аминопропановая кислота","аланин (α-аланин)"),
]


def set_font(run, bold=False, size=12):
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def build_document():
    doc = Document()

    # Поля страницы
    section = doc.sections[0]
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(1.5)

    # Стиль Normal → Times New Roman 12
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    # Заголовок
    h = doc.add_paragraph()
    r = h.add_run("Тренажер по номенклатуре органических соединений")
    set_font(r, bold=True, size=13)

    doc.add_paragraph()  # пустая строка

    # Задание
    task_p = doc.add_paragraph()
    task_r = task_p.add_run("Задание: ")
    set_font(task_r, bold=True)
    rest_r = task_p.add_run(
        "определить систематические названия соединений и, где возможно, "
        "указать тривиальные названия."
    )
    set_font(rest_r)

    doc.add_paragraph()

    # Ответы — простой нумерованный список
    for num, formula, systematic, trivial in ANSWERS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)

        # Номер
        r_num = p.add_run(f"{num}. ")
        set_font(r_num, bold=True)

        # Формула
        r_formula = p.add_run(f"{formula}  —  ")
        set_font(r_formula)

        # Систематическое
        r_sys = p.add_run(systematic)
        set_font(r_sys, bold=True)

        # Тривиальное (если есть)
        if trivial:
            r_triv = p.add_run(f"  ({trivial})")
            set_font(r_triv)

    doc.save(OUTPUT_DOCX)
    print(f"Готово: {OUTPUT_DOCX}")


if __name__ == "__main__":
    build_document()
