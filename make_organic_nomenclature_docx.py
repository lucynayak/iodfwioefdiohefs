from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


OUTPUT_DOCX = "Тренажер по номенклатуре органических соединений - ОТВЕТЫ.docx"


SECTIONS = [
    (
        "Простейшие углеводороды",
        [
            (1, "CH3-CH2-CH3", "Пропан", "—"),
            (2, "CH3-CH2-CH2-CH3", "Бутан", "н-бутан"),
            (3, "CH2=CH2", "Этен", "этилен"),
            (4, "HC≡CH", "Этин", "ацетилен"),
            (5, "C6H6 (бензольное кольцо)", "Бензол", "—"),
        ],
    ),
    (
        "Спирты и фенолы",
        [
            (6, "CH3OH", "Метанол", "метиловый спирт"),
            (7, "CH3-CH2-OH", "Этанол", "этиловый спирт"),
            (8, "CH3-CH(OH)-CH3", "Пропан-2-ол", "изопропиловый спирт, изопропанол"),
            (9, "HO-CH2-CH2-OH", "Этан-1,2-диол", "этиленгликоль"),
            (10, "C6H5OH", "Гидроксибензол", "фенол"),
        ],
    ),
    (
        "Альдегиды и кетоны",
        [
            (11, "HCHO", "Метаналь", "формальдегид"),
            (12, "CH3-CHO", "Этаналь", "ацетальдегид"),
            (13, "CH3-CO-CH3", "Пропанон", "ацетон"),
            (14, "CH3-CH2-CHO", "Пропаналь", "пропионовый альдегид"),
        ],
    ),
    (
        "Карбоновые кислоты",
        [
            (15, "HCOOH", "Метановая кислота", "муравьиная кислота"),
            (16, "CH3COOH", "Этановая кислота", "уксусная кислота"),
            (17, "CH3-CH2-COOH", "Пропановая кислота", "пропионовая кислота"),
            (18, "HOOC-COOH", "Этандиовая кислота", "щавелевая кислота"),
        ],
    ),
    (
        "Амины и нитросоединения",
        [
            (19, "CH3-NH2", "Метанамин", "метиламин"),
            (20, "CH3-NO2", "Нитрометан", "—"),
        ],
    ),
    (
        "Эфиры и галогенпроизводные",
        [
            (21, "CH3-O-CH3", "Метоксиметан", "диметиловый эфир"),
            (22, "CH3-COO-CH3", "Метилэтаноат", "метилацетат"),
            (23, "CH3-CH2-Cl", "Хлорэтан", "этилхлорид"),
            (24, "Cl-CH2-CH2-Cl", "1,2-Дихлорэтан", "этилендихлорид"),
        ],
    ),
    (
        "Разветвленные цепи",
        [
            (25, "C(CH3)4", "2,2-Диметилпропан", "неопентан"),
            (26, "CH3-CH(CH3)-CH2-CH3", "2-Метилбутан", "изопентан"),
            (27, "CH3-C(CH3)=CH-CH3", "2-Метилбут-2-ен", "изоамилен"),
        ],
    ),
    (
        "Циклические соединения",
        [
            (28, "(CH2)5 (циклическая цепь)", "Циклопентан", "—"),
            (29, "C6H5-CH3", "Метилбензол", "толуол"),
            (30, "CH3-CH(NH2)-COOH", "2-Аминопропановая кислота", "аланин"),
        ],
    ),
]


def set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)


def format_header_row(cells) -> None:
    headers = [
        "№",
        "Формула",
        "Систематическое название",
        "Тривиальное / распространенное название",
    ]
    for cell, value in zip(cells, headers):
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(value)
        run.bold = True


def add_section_table(document: Document, title: str, rows_data) -> None:
    heading = document.add_paragraph()
    heading.style = document.styles["Heading 2"]
    heading.add_run(title)

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    format_header_row(table.rows[0].cells)

    for number, formula, systematic, trivial in rows_data:
        cells = table.add_row().cells
        cells[0].text = str(number)
        cells[1].text = formula
        cells[2].text = systematic
        cells[3].text = trivial

    document.add_paragraph()


def build_document() -> None:
    document = Document()
    set_default_font(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Тренажер по номенклатуре органических соединений")
    title_run.bold = True
    title_run.font.size = Pt(14)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Ответы: систематические и тривиальные названия")

    note = document.add_paragraph()
    note.add_run(
        "Примечание: в школьной записи могут встречаться равноправные варианты вроде "
        "«пропан-2-ол / 2-пропанол» или «метилэтаноат / метилацетат»."
    )

    for section_title, rows_data in SECTIONS:
        add_section_table(document, section_title, rows_data)

    document.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
    print(f"Документ создан: {OUTPUT_DOCX}")